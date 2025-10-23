"""
Async RAG Engine - Core business logic without CLI dependencies
"""

import os
import asyncio
import logging
from pathlib import Path
from typing import List, Optional, Dict, Any, Callable
import hashlib
import json
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime
import tempfile
import time

import pypdf
import pdfplumber
from docx import Document as DocxDocument

# 可选的文件类型检测
try:
    import magic
    MAGIC_AVAILABLE = True
except (ImportError, OSError):
    MAGIC_AVAILABLE = False

from llama_index.core import (
    Document,
    Settings,
    StorageContext,
    VectorStoreIndex,
)
from llama_index.core.node_parser import SentenceSplitter
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.llms.openai_like import OpenAILike

from elasticsearch import Elasticsearch
from elasticsearch.helpers import bulk

try:
    from llama_index.vector_stores.elasticsearch import ElasticsearchVectorStore
except ImportError:  # pragma: no cover - fallback for older LlamaIndex versions
    from llama_index.vector_stores.elasticsearch import ElasticsearchStore as ElasticsearchVectorStore


class AsyncRAGEngine:
    """异步RAG引擎 - 去除CLI依赖，支持Web API调用"""
    
    def __init__(
        self,
        data_dir: str = "data",
        persist_dir: str = "storage",
        embedding_model: str = "BAAI/bge-small-zh-v1.5",
        enable_parallel: bool = True,
        max_workers: int = 2,
        es_url: Optional[str] = None,
        es_index: str = "kb-documents",
        es_keyword_index: str = "kb-documents-text",
        es_user: Optional[str] = None,
        es_password: Optional[str] = None,
        es_text_analyzer: Optional[str] = None,
    ):
        """初始化异步RAG引擎"""
        self.data_dir = Path(data_dir)
        self.persist_dir = Path(persist_dir)
        self.index = None
        self.logger = self._setup_logger()
        
        # 性能优化配置
        self.enable_parallel = enable_parallel
        self.max_workers = max_workers
        self.cache_dir = self.persist_dir / "cache"
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        
        # 文件处理缓存
        self.file_cache_path = self.cache_dir / "file_cache.json"
        self.file_cache = self._load_file_cache()
        self.vector_store = None
        self.storage_context = None
        self.es_client = None
        
        # 配置LlamaIndex设置 - 使用自定义DeepSeek LLM
        Settings.llm = OpenAILike(
            model="deepseek-v3-250324",
            api_key="155d5cb5-6b83-4d52-8be8-eb795c72ad44",
            api_base="https://ark.cn-beijing.volces.com/api/v3",
            is_chat_model=True,
            temperature=0.1
        )
        
        # 使用本地embedding模型
        self._setup_embedding_model(embedding_model)
        Settings.node_parser = SentenceSplitter(chunk_size=1024, chunk_overlap=20)
        
        # Elasticsearch 相关配置
        self.es_url = es_url or os.getenv("ELASTICSEARCH_URL", "http://localhost:9200")
        self.es_index = os.getenv("ELASTICSEARCH_INDEX", es_index)
        self.es_keyword_index = os.getenv("ELASTICSEARCH_TEXT_INDEX", es_keyword_index)
        self.es_user = es_user or os.getenv("ELASTICSEARCH_USER")
        self.es_password = es_password or os.getenv("ELASTICSEARCH_PASSWORD")
        self.es_api_key = os.getenv("ELASTICSEARCH_API_KEY")
        verify_env = os.getenv("ELASTICSEARCH_VERIFY_CERTS")
        self.es_verify_certs: Optional[bool] = None
        if verify_env is not None:
            self.es_verify_certs = verify_env.lower() not in {"false", "0", "no"}
        self.es_ca_certs = os.getenv("ELASTICSEARCH_CA_CERTS")
        timeout_env = os.getenv("ELASTICSEARCH_TIMEOUT")
        self.es_timeout = int(timeout_env) if timeout_env and timeout_env.isdigit() else None
        analyzer_env = os.getenv("ELASTICSEARCH_TEXT_ANALYZER")
        self.es_text_analyzer = analyzer_env or es_text_analyzer or "standard"
        self._keyword_index_checked = False

        # 线程池用于并行处理
        self._thread_pool = ThreadPoolExecutor(max_workers=max_workers)

    def _get_es_kwargs(self) -> Dict[str, Any]:
        """构建 Elasticsearch 连接参数"""
        es_kwargs: Dict[str, Any] = {}
        if self.es_user and self.es_password:
            es_kwargs["basic_auth"] = (self.es_user, self.es_password)
        if self.es_api_key:
            es_kwargs["api_key"] = self.es_api_key
        if self.es_verify_certs is not None:
            es_kwargs["verify_certs"] = self.es_verify_certs
        if self.es_ca_certs:
            es_kwargs["ca_certs"] = self.es_ca_certs
        if self.es_timeout:
            es_kwargs["request_timeout"] = self.es_timeout
        return es_kwargs

    def _ensure_es_client(self, force: bool = False):
        """确保 Elasticsearch 客户端已初始化"""
        if self.es_client is not None and not force:
            return

        es_kwargs = self._get_es_kwargs()

        try:
            self.es_client = Elasticsearch(self.es_url, **es_kwargs)
        except Exception as exc:
            self.logger.error(f"连接 Elasticsearch 失败: {exc}")
            raise

        try:
            self.es_client.ping()
        except Exception as exc:  # pragma: no cover - ping 失败不阻断
            self.logger.warning(f"Elasticsearch ping 失败: {exc}")

    def _ensure_keyword_index(self, force: bool = False):
        """确保关键字检索索引已准备"""
        self._ensure_es_client()

        if force:
            try:
                self.es_client.indices.delete(index=self.es_keyword_index, ignore_unavailable=True)
                self.logger.info(f"已删除 Elasticsearch 关键字索引 {self.es_keyword_index}")
            except Exception as exc:
                self.logger.warning(f"删除关键字索引失败: {exc}")
            finally:
                self._keyword_index_checked = False

        if self._keyword_index_checked:
            return

        try:
            exists = self.es_client.indices.exists(index=self.es_keyword_index)
        except Exception as exc:
            self.logger.error(f"检测关键字索引失败: {exc}")
            raise

        if not exists:
            index_body = {
                "mappings": {
                    "properties": {
                        "text": {"type": "text", "analyzer": self.es_text_analyzer},
                        "filename": {"type": "keyword"},
                        "file_path": {"type": "keyword"},
                        "file_type": {"type": "keyword"},
                        "document_id": {"type": "keyword"},
                        "chunk_id": {"type": "keyword"},
                        "position": {"type": "integer"},
                        "char_count": {"type": "integer"},
                        "metadata": {"type": "object", "enabled": False},
                        "created_at": {"type": "date", "format": "strict_date_optional_time||epoch_millis"},
                    }
                }
            }

            try:
                self.es_client.indices.create(index=self.es_keyword_index, body=index_body)
                self.logger.info(f"Elasticsearch 关键字索引已创建: {self.es_keyword_index}")
            except Exception as exc:
                self.logger.error(f"创建关键字索引失败: {exc}")
                raise

        self._keyword_index_checked = True

    def _reset_keyword_index(self):
        """强制重建关键字索引"""
        self._ensure_keyword_index(force=True)

    def _setup_logger(self) -> logging.Logger:
        """配置日志记录器"""
        logger = logging.getLogger(f"{__name__}.AsyncRAGEngine")
        logger.setLevel(logging.INFO)
        
        if not logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter(
                '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
            )
            handler.setFormatter(formatter)
            logger.addHandler(handler)
            
        return logger
        
    def _setup_embedding_model(self, model_name: str):
        """配置embedding模型"""
        self.logger.info(f"正在加载embedding模型: {model_name}")
        
        # 预定义的模型映射
        model_info = {
            "BAAI/bge-small-zh-v1.5": "BGE小型中文模型 (推荐中文使用)",
            "BAAI/bge-base-zh-v1.5": "BGE基础中文模型 (更好效果但更大)",
            "sentence-transformers/all-MiniLM-L6-v2": "轻量级英文模型 (快速)",
            "sentence-transformers/all-mpnet-base-v2": "高质量英文模型",
            "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2": "多语言模型",
        }
        
        if model_name in model_info:
            self.logger.info(f"使用模型: {model_info[model_name]}")
        
        try:
            Settings.embed_model = HuggingFaceEmbedding(model_name=model_name)
            self.logger.info("Embedding模型加载成功")
        except Exception as e:
            self.logger.error(f"加载embedding模型失败: {e}")
            self.logger.info("尝试使用默认的轻量级模型...")
            Settings.embed_model = HuggingFaceEmbedding(
                model_name="sentence-transformers/all-MiniLM-L6-v2"
            )
    
    def _init_vector_store(self, force: bool = False):
        """初始化或重建 Elasticsearch 向量存储"""
        if self.vector_store is not None and self.storage_context is not None and not force:
            return
        
        self._ensure_es_client(force=force)

        try:
            self.vector_store = ElasticsearchVectorStore(
                index_name=self.es_index,
                es_url=self.es_url,
                es_user=self.es_user,
                es_password=self.es_password,
                es_api_key=self.es_api_key,
            )
            self.storage_context = StorageContext.from_defaults(vector_store=self.vector_store)
            self.logger.info(f"Elasticsearch 向量索引已准备: {self.es_index}")
        except Exception as exc:
            self.logger.error(f"初始化 Elasticsearch 向量存储失败: {exc}")
            raise
    
    def _ensure_vector_store(self):
        """确保向量存储已初始化"""
        if self.vector_store is None or self.storage_context is None:
            self._init_vector_store()
    
    def _reset_vector_index(self):
        """删除并重建 Elasticsearch 索引"""
        self._ensure_vector_store()
        
        try:
            delete_method = getattr(self.vector_store, "delete_index", None)
            if callable(delete_method):
                delete_method()
                self.logger.info(f"已清空 Elasticsearch 向量索引 {self.es_index}")
            else:
                self.es_client.indices.delete(index=self.es_index, ignore_unavailable=True)
                self.logger.info(f"已删除 Elasticsearch 索引 {self.es_index}")
        except Exception as exc:
            self.logger.warning(f"删除 Elasticsearch 索引失败: {exc}")
        finally:
            self._init_vector_store(force=True)
    
    def _refresh_vector_index(self):
        """刷新 Elasticsearch 索引以便查询"""
        if not self.es_client:
            return
        try:
            self.es_client.indices.refresh(index=self.es_index)
        except Exception as exc:
            self.logger.warning(f"刷新 Elasticsearch 索引失败: {exc}")
    
    def _load_file_cache(self) -> Dict[str, Any]:
        """加载文件处理缓存"""
        try:
            if self.file_cache_path.exists():
                with open(self.file_cache_path, 'r', encoding='utf-8') as f:
                    return json.load(f)
        except Exception as e:
            self.logger.warning(f"缓存加载失败: {e}")
        return {}
    
    def _save_file_cache(self):
        """保存文件处理缓存"""
        try:
            with open(self.file_cache_path, 'w', encoding='utf-8') as f:
                json.dump(self.file_cache, f, ensure_ascii=False, indent=2)
        except Exception as e:
            self.logger.warning(f"缓存保存失败: {e}")
    
    def _get_file_hash(self, file_path: Path) -> str:
        """计算文件的哈希值用于缓存判断"""
        try:
            stat = file_path.stat()
            content = f"{file_path.name}_{stat.st_size}_{stat.st_mtime}"
            return hashlib.md5(content.encode()).hexdigest()
        except Exception:
            return ""
    
    def _is_file_cached(self, file_path: Path) -> bool:
        """检查文件是否已被缓存且未过期"""
        file_hash = self._get_file_hash(file_path)
        if not file_hash:
            return False
        
        cache_key = str(file_path)
        if cache_key in self.file_cache:
            cached_info = self.file_cache[cache_key]
            return cached_info.get('hash') == file_hash
        return False
    
    def _get_cached_content(self, file_path: Path) -> Optional[str]:
        """获取缓存的文件内容"""
        cache_key = str(file_path)
        if cache_key in self.file_cache:
            return self.file_cache[cache_key].get('content')
        return None
    
    def _cache_file_content(self, file_path: Path, content: str):
        """缓存文件内容"""
        file_hash = self._get_file_hash(file_path)
        if file_hash:
            cache_key = str(file_path)
            self.file_cache[cache_key] = {
                'hash': file_hash,
                'content': content,
                'timestamp': time.time(),
                'char_count': len(content)
            }
    
    def _detect_file_type(self, file_path: Path) -> str:
        """检测文件的真实类型"""
        if MAGIC_AVAILABLE:
            try:
                mime_type = magic.from_file(str(file_path), mime=True)
                if mime_type == 'application/pdf':
                    return 'pdf'
                elif mime_type in {
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'application/msword',
                }:
                    return 'docx'
                elif mime_type.startswith('text/'):
                    return 'txt'
                else:
                    return file_path.suffix.lower().lstrip('.')
            except Exception:
                return file_path.suffix.lower().lstrip('.')
        else:
            suffix = file_path.suffix.lower().lstrip('.')
            if suffix in {'docx', 'doc'}:
                return 'docx'
            return suffix
    
    def _extract_pdf_content(self, file_path: Path) -> Optional[str]:
        """从PDF文件提取文本内容"""
        # 首先检查缓存
        if self._is_file_cached(file_path):
            cached_content = self._get_cached_content(file_path)
            if cached_content:
                self.logger.info(f"使用缓存内容: {file_path.name} ({len(cached_content)} 字符)")
                return cached_content
        
        # 检查文件大小和基本有效性
        try:
            file_size = file_path.stat().st_size
            if file_size == 0:
                self.logger.warning(f"跳过空文件: {file_path.name}")
                return None
                
            if file_size > 50 * 1024 * 1024:  # 50MB
                self.logger.warning(f"大文件 {file_path.name} ({file_size / 1024 / 1024:.1f}MB)，处理可能较慢")
                
        except OSError as e:
            self.logger.error(f"无法访问文件 {file_path.name}: {e}")
            return None
        
        try:
            # 首先尝试使用pypdf
            self.logger.info(f"使用pypdf提取 {file_path.name}")
            text = self._extract_with_pypdf(file_path)
            
            if text and text.strip():
                self.logger.info(f"pypdf成功提取 {len(text)} 字符")
                self._cache_file_content(file_path, text)
                return text
            
            # 如果pypdf失败，尝试pdfplumber
            self.logger.warning(f"pypdf提取结果为空，尝试pdfplumber: {file_path.name}")
            text = self._extract_with_pdfplumber(file_path)
            
            if text and text.strip():
                self.logger.info(f"pdfplumber成功提取 {len(text)} 字符")
                self._cache_file_content(file_path, text)
                return text
            else:
                self.logger.warning(f"PDF文件 {file_path.name} 可能是扫描版或损坏")
                return None
            
        except Exception as e:
            error_type = type(e).__name__
            self.logger.error(f"PDF提取失败 {file_path.name} ({error_type}): {str(e)[:100]}")
            return None

    def _extract_docx_content(self, file_path: Path) -> Optional[str]:
        """从 DOCX 文件提取文本内容"""
        try:
            doc = DocxDocument(str(file_path))
        except Exception as exc:
            self.logger.error(f"DOCX 读取失败 {file_path.name}: {exc}")
            return None

        paragraphs: List[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)

        # 处理表格中的文本
        for table in getattr(doc, "tables", []):
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        paragraphs.append(cell_text)

        content = "\n".join(paragraphs).strip()
        if content:
            self._cache_file_content(file_path, content)
            return content

        self.logger.warning(f"DOCX文件 {file_path.name} 未提取到有效文本")
        return None
    
    def _extract_with_pypdf(self, file_path: Path) -> str:
        """使用pypdf提取PDF文本"""
        text = ""
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                if total_pages == 0:
                    self.logger.warning(f"PDF文件 {file_path.name} 没有页面")
                    return ""
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += f"\n--- 第{page_num + 1}页 ---\n{page_text}\n"
                    except Exception as e:
                        self.logger.warning(f"跳过第{page_num + 1}页: {str(e)[:50]}")
                        continue
                        
        except Exception as e:
            raise Exception(f"pypdf读取失败: {e}")
            
        return text.strip()
    
    def _extract_with_pdfplumber(self, file_path: Path) -> str:
        """使用pdfplumber提取PDF文本"""
        text = ""
        
        try:
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                
                if total_pages == 0:
                    self.logger.warning(f"PDF文件 {file_path.name} 没有页面")
                    return ""
                
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += f"\n--- 第{page_num + 1}页 ---\n{page_text}\n"
                    except Exception as e:
                        self.logger.warning(f"跳过第{page_num + 1}页: {str(e)[:50]}")
                        continue
                        
        except Exception as e:
            raise Exception(f"pdfplumber处理失败: {e}")
            
        return text.strip()
    
    def _process_single_file(self, file_path: Path) -> Optional[Document]:
        """处理单个文件，返回Document对象或None"""
        detected_type = self._detect_file_type(file_path)
        
        try:
            if detected_type == 'pdf':
                content = self._extract_pdf_content(file_path)
                if content:
                    return Document(
                        text=content, 
                        metadata={
                            "filename": file_path.name,
                            "file_type": "pdf",
                            "file_path": str(file_path),
                            "file_size": file_path.stat().st_size,
                            "char_count": len(content)
                        }
                    )
                    
            elif detected_type == 'docx':
                content = self._extract_docx_content(file_path)
                if content:
                    return Document(
                        text=content,
                        metadata={
                            "filename": file_path.name,
                            "file_type": "docx",
                            "file_path": str(file_path),
                            "file_size": file_path.stat().st_size,
                            "char_count": len(content),
                        },
                    )

            elif detected_type == 'txt':
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                        if content.strip():
                            return Document(
                                text=content, 
                                metadata={
                                    "filename": file_path.name,
                                    "file_type": "txt",
                                    "file_path": str(file_path),
                                    "file_size": file_path.stat().st_size,
                                    "char_count": len(content)
                                }
                            )
                except UnicodeDecodeError:
                    self.logger.warning(f"文本文件编码错误，跳过: {file_path.name}")
                    
        except Exception as e:
            self.logger.error(f"处理文件失败 {file_path.name}: {str(e)[:100]}")
        
        return None
    
    async def process_documents_async(self, 
                                    files: List[Path], 
                                    progress_callback: Optional[Callable] = None) -> List[Document]:
        """异步处理文档列表"""
        documents = []
        stats = {
            'processed': 0,
            'failed': 0,
            'total_chars': 0,
            'pdf_files': 0,
            'txt_files': 0
        }
        
        start_time = time.time()
        
        if self.enable_parallel and len(files) > 1:
            documents = await self._process_files_parallel_async(files, stats, progress_callback)
        else:
            documents = await self._process_files_sequential_async(files, stats, progress_callback)
        
        # 保存缓存
        self._save_file_cache()
        
        processing_time = time.time() - start_time
        self.logger.info(f"处理完成: {stats['processed']} 个文档, 用时 {processing_time:.2f}秒")
        
        return documents

    def _create_nodes(self, documents: List[Document]):
        """根据配置的分句器将文档切分为节点"""
        parser = Settings.node_parser
        if parser is None:
            parser = SentenceSplitter(chunk_size=1024, chunk_overlap=20)
        return parser.get_nodes_from_documents(documents)

    def _document_from_bytes(self, filename: str, data: bytes) -> Optional[Document]:
        """将上传的文件字节转换为 Document"""
        suffix = Path(filename or "").suffix.lower()
        allowed_suffixes = {".pdf", ".txt", ".docx"}
        if suffix not in allowed_suffixes:
            raise ValueError(f"不支持的文件类型: {suffix or 'unknown'}")

        if suffix == ".txt":
            try:
                text = data.decode("utf-8")
            except UnicodeDecodeError:
                text = data.decode("utf-8", errors="ignore")

            text = text.strip()
            if not text:
                return None

            metadata = {
                "filename": filename or "uploaded.txt",
                "file_type": "txt",
                "file_path": f"upload://{filename}" if filename else None,
                "file_size": len(data),
                "char_count": len(text),
            }
            metadata = {k: v for k, v in metadata.items() if v is not None}
            return Document(text=text, metadata=metadata)

        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix or ".bin") as tmp_file:
            tmp_file.write(data)
            temp_path = Path(tmp_file.name)

        try:
            doc = self._process_single_file(temp_path)
            if doc:
                metadata = dict(doc.metadata or {})
                metadata.setdefault("filename", filename or temp_path.name)
                metadata["file_path"] = f"upload://{filename}" if filename else metadata.get("file_path")
                metadata.setdefault("file_size", len(data))
                metadata.setdefault("char_count", len(doc.text))
                doc.metadata = metadata
            return doc
        finally:
            try:
                temp_path.unlink()
            except OSError:
                pass

    async def documents_from_uploads(self, uploads: List[Dict[str, Any]]) -> List[Document]:
        """将上传的文件数据转换为 Document 列表"""
        loop = asyncio.get_event_loop()
        tasks = []
        for item in uploads:
            filename = item.get("filename", "")
            data = item.get("data")
            if not isinstance(data, (bytes, bytearray)):
                raise ValueError("文件数据必须是字节格式")
            tasks.append(
                loop.run_in_executor(
                    self._thread_pool,
                    self._document_from_bytes,
                    filename,
                    bytes(data),
                )
            )

        documents: List[Document] = []
        for task in tasks:
            doc = await task
            if doc:
                documents.append(doc)
        return documents

    def _index_keyword_nodes(self, nodes: List[Any]) -> int:
        """将切分后的文本节点写入关键字索引，并返回写入数量"""
        if not nodes:
            return 0

        self._ensure_keyword_index()

        actions = []
        chunk_counters: Dict[str, int] = {}

        for node in nodes:
            # TextNode 在 LlamaIndex 中既有 text 字段，也可以通过 get_content 获取
            text = getattr(node, "text", None) or node.get_content(metadata_mode="none")
            if not text or not text.strip():
                continue

            metadata = getattr(node, "metadata", {}) or {}
            document_key = (
                metadata.get("file_path")
                or metadata.get("filename")
                or getattr(node, "ref_doc_id", None)
                or "unknown_document"
            )
            doc_hash = hashlib.md5(document_key.encode("utf-8")).hexdigest()
            position = chunk_counters.get(doc_hash, 0)
            chunk_counters[doc_hash] = position + 1
            chunk_id = f"{doc_hash}-{position}"

            action = {
                "_op_type": "index",
                "_index": self.es_keyword_index,
                "_id": chunk_id,
                "document_id": doc_hash,
                "chunk_id": chunk_id,
                "filename": metadata.get("filename"),
                "file_path": metadata.get("file_path"),
                "file_type": metadata.get("file_type"),
                "position": position,
                "char_count": len(text),
                "text": text,
                "metadata": metadata,
                "created_at": datetime.utcnow().isoformat(),
            }

            actions.append(action)

        if not actions:
            return 0

        try:
            bulk(
                self.es_client,
                actions,
                chunk_size=500,
                request_timeout=self.es_timeout or 30,
            )
            self.es_client.indices.refresh(index=self.es_keyword_index)
            self.logger.info(
                f"关键字索引 {self.es_keyword_index} 已更新 {len(actions)} 个文本分块"
            )
            return len(actions)
        except Exception as exc:
            self.logger.error(f"写入关键字索引失败: {exc}")
            raise
        return len(actions)

    async def _process_files_parallel_async(self, 
                                          files: List[Path], 
                                          stats: Dict[str, int],
                                          progress_callback: Optional[Callable] = None) -> List[Document]:
        """异步并行处理文件"""
        documents = []
        
        self.logger.info(f"使用并行处理 (最大 {self.max_workers} 个工作线程)")
        
        loop = asyncio.get_event_loop()
        
        # 创建任务
        tasks = []
        for file_path in files:
            task = loop.run_in_executor(
                self._thread_pool, 
                self._process_single_file, 
                file_path
            )
            tasks.append((task, file_path))
        
        # 收集结果
        for task, file_path in tasks:
            try:
                doc = await task
                if doc:
                    documents.append(doc)
                    file_type = doc.metadata.get('file_type', 'unknown')
                    if file_type == 'pdf':
                        stats['pdf_files'] += 1
                    elif file_type == 'txt':
                        stats['txt_files'] += 1
                    
                    stats['processed'] += 1
                    stats['total_chars'] += doc.metadata.get('char_count', 0)
                else:
                    stats['failed'] += 1
                
                # 调用进度回调
                if progress_callback:
                    await progress_callback({
                        'file': file_path.name,
                        'status': 'completed' if doc else 'failed',
                        'processed': stats['processed'],
                        'total': len(files)
                    })
                    
            except Exception as e:
                self.logger.error(f"并行处理失败 {file_path.name}: {e}")
                stats['failed'] += 1
        
        return documents
    
    async def _process_files_sequential_async(self, 
                                            files: List[Path], 
                                            stats: Dict[str, int],
                                            progress_callback: Optional[Callable] = None) -> List[Document]:
        """异步串行处理文件"""
        documents = []
        
        loop = asyncio.get_event_loop()
        
        for file_path in files:
            try:
                doc = await loop.run_in_executor(
                    self._thread_pool, 
                    self._process_single_file, 
                    file_path
                )
                
                if doc:
                    documents.append(doc)
                    file_type = doc.metadata.get('file_type', 'unknown')
                    if file_type == 'pdf':
                        stats['pdf_files'] += 1
                    elif file_type == 'txt':
                        stats['txt_files'] += 1
                    
                    stats['processed'] += 1
                    stats['total_chars'] += doc.metadata.get('char_count', 0)
                else:
                    stats['failed'] += 1
                
                # 调用进度回调
                if progress_callback:
                    await progress_callback({
                        'file': file_path.name,
                        'status': 'completed' if doc else 'failed',
                        'processed': stats['processed'],
                        'total': len(files)
                    })
                    
            except Exception as e:
                self.logger.error(f"处理文件失败 {file_path.name}: {e}")
                stats['failed'] += 1
        
        return documents
    
    async def build_index_async(
        self,
        documents: Optional[List[Document]] = None,
        reset_existing: bool = False,
    ) -> VectorStoreIndex:
        """异步构建向量索引（写入 Elasticsearch）"""
        self.logger.info("构建向量索引（Elasticsearch）...")
        self._ensure_vector_store()
        
        if documents is None:
            # 加载数据目录中的文档
            all_files = [f for f in self.data_dir.iterdir() if f.is_file()]
            documents = await self.process_documents_async(all_files)
        
        if not documents:
            raise ValueError("没有找到要索引的文档")

        loop = asyncio.get_event_loop()

        nodes = self._create_nodes(documents)
        if not nodes:
            raise ValueError("无法从文档生成有效的文本分块")

        if reset_existing:
            await loop.run_in_executor(None, self._reset_vector_index)
            await loop.run_in_executor(None, self._reset_keyword_index)

        def _build() -> VectorStoreIndex:
            return VectorStoreIndex.from_nodes(
                nodes,
                storage_context=self.storage_context,
            )

        index = await loop.run_in_executor(None, _build)
        await loop.run_in_executor(None, self._refresh_vector_index)
        await loop.run_in_executor(None, lambda: self._index_keyword_nodes(nodes))

        self.logger.info("Elasticsearch 向量索引构建完成")
        self.index = index
        return index
    
    async def ingest_documents_async(
        self,
        documents: List[Dict[str, Any]],
        reset_existing: bool = False,
    ) -> Dict[str, int]:
        """接收外部服务提交的文档并写入索引"""
        if not documents:
            raise ValueError("文档列表不能为空")

        self._ensure_vector_store()
        loop = asyncio.get_event_loop()

        doc_objects: List[Document] = []
        raw_chunks = 0

        for item in documents:
            document_id = item.get("document_id") or item.get("id")
            if not document_id:
                self.logger.warning("跳过缺少 document_id 的文档: %s", item)
                continue

            base_metadata: Dict[str, Any] = dict(item.get("metadata") or {})
            base_metadata.setdefault("document_id", document_id)
            filename = item.get("filename")
            if filename:
                base_metadata.setdefault("filename", filename)
            base_metadata.setdefault("file_type", item.get("file_type", "external"))
            base_metadata.setdefault("source", item.get("source", "remote_ingest"))

            chunk_list = item.get("chunks") or []
            if chunk_list:
                for idx, chunk in enumerate(chunk_list):
                    text = (chunk.get("content") or "").strip()
                    if not text:
                        continue
                    chunk_metadata = base_metadata.copy()
                    chunk_metadata.update(chunk.get("metadata") or {})
                    chunk_metadata["chunk_index"] = idx
                    chunk_metadata["char_count"] = len(text)
                    doc_objects.append(Document(text=text, metadata=chunk_metadata))
                    raw_chunks += 1
            else:
                text = (item.get("content") or "").strip()
                if not text:
                    continue
                chunk_metadata = base_metadata.copy()
                chunk_metadata["chunk_index"] = 0
                chunk_metadata["char_count"] = len(text)
                doc_objects.append(Document(text=text, metadata=chunk_metadata))
                raw_chunks += 1

        if not doc_objects:
            raise ValueError("没有可供索引的文档内容")

        nodes = self._create_nodes(doc_objects)

        created_new_index = False

        if reset_existing:
            await loop.run_in_executor(None, self._reset_vector_index)
            await loop.run_in_executor(None, self._reset_keyword_index)
            self.index = None

        if self.index is None and not reset_existing:
            existing_index = await self.load_index_async()
            if existing_index is not None:
                self.index = existing_index

        if self.index is None:
            def _build_nodes() -> VectorStoreIndex:
                return VectorStoreIndex.from_nodes(
                    nodes,
                    storage_context=self.storage_context,
                )

            self.index = await loop.run_in_executor(None, _build_nodes)
            created_new_index = True
        else:
            await loop.run_in_executor(None, lambda: self.index.insert_nodes(nodes))

        await loop.run_in_executor(None, self._refresh_vector_index)
        keyword_chunks = await loop.run_in_executor(
            None,
            lambda: self._index_keyword_nodes(nodes),
        )

        return {
            "indexed_documents": len(documents),
            "indexed_nodes": len(nodes),
            "keyword_chunks": int(keyword_chunks or 0),
            "raw_chunks": raw_chunks,
            "created_new_index": created_new_index,
        }
    
    async def load_index_async(self) -> Optional[VectorStoreIndex]:
        """异步加载已存在的 Elasticsearch 索引"""
        self._ensure_vector_store()
        loop = asyncio.get_event_loop()
        
        try:
            index_exists = await loop.run_in_executor(
                None,
                lambda: self.es_client.indices.exists(index=self.es_index),
            )
        except Exception as exc:
            self.logger.warning(f"检测 Elasticsearch 索引失败: {exc}")
            return None
        
        if not index_exists:
            self.logger.info("Elasticsearch 索引不存在，需要重新构建")
            return None
        
        def _load() -> VectorStoreIndex:
            return VectorStoreIndex.from_vector_store(
                self.vector_store,
                storage_context=self.storage_context,
            )
        
        try:
            index = await loop.run_in_executor(None, _load)
            self.logger.info("成功连接到 Elasticsearch 向量索引")
            self.index = index
            return index
        except Exception as exc:
            self.logger.warning(f"加载 Elasticsearch 索引失败: {exc}")
            return None
    
    async def get_or_create_index_async(self) -> VectorStoreIndex:
        """异步获取或创建索引"""
        index = await self.load_index_async()
        if index is None:
            index = await self.build_index_async()
        
        self.index = index
        return index
    
    async def query_async(
        self,
        question: str,
        search_params: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """根据请求参数执行 RAG 或关键字检索"""
        params = search_params or {}
        mode = str(params.get("mode", "rag")).lower()

        if mode == "es":
            top_k = params.get("top_k") or params.get("keyword_top_k") or 5
            min_score = params.get("min_score")
            return await self.keyword_search_async(
                question,
                top_k=int(top_k) if isinstance(top_k, (int, float, str)) else 5,
                min_score=min_score,
            )

        top_k = params.get("top_k") or params.get("similarity_top_k")
        response_mode = params.get("response_mode", "compact")
        return await self.rag_search_async(
            question,
            similarity_top_k=top_k,
            response_mode=response_mode,
        )

    async def rag_search_async(
        self,
        question: str,
        similarity_top_k: Optional[Any] = None,
        response_mode: str = "compact",
    ) -> Dict[str, Any]:
        """执行基于向量的 RAG 检索"""
        self._ensure_vector_store()
        if self.index is None:
            await self.get_or_create_index_async()

        loop = asyncio.get_event_loop()
        top_k = 5
        if similarity_top_k is not None:
            try:
                top_k = max(1, int(similarity_top_k))
            except (TypeError, ValueError):
                top_k = 5

        query_engine = self.index.as_query_engine(
            response_mode=response_mode,
            similarity_top_k=top_k,
        )

        response = await loop.run_in_executor(
            None,
            query_engine.query,
            question,
        )

        sources = self._format_sources_from_nodes(response)

        return {
            "answer": str(response),
            "sources": sources,
        }

    async def keyword_search_async(
        self,
        question: str,
        top_k: int = 5,
        min_score: Optional[Any] = None,
    ) -> Dict[str, Any]:
        """执行基于关键字的 Elasticsearch 检索"""
        self._ensure_es_client()
        self._ensure_keyword_index()

        loop = asyncio.get_event_loop()

        try:
            top_k_value = max(1, int(top_k))
        except (TypeError, ValueError):
            top_k_value = 5

        min_score_value: Optional[float] = None
        if min_score is not None:
            try:
                min_score_value = float(min_score)
            except (TypeError, ValueError):
                min_score_value = None

        def _search():
            body = {
                "size": top_k_value,
                "query": {
                    "multi_match": {
                        "query": question,
                        "fields": ["text^3", "filename"],
                        "type": "best_fields",
                    }
                },
                "highlight": {
                    "fields": {
                        "text": {
                            "fragment_size": 120,
                            "number_of_fragments": 1,
                        }
                    }
                },
            }
            if min_score_value is not None:
                body["min_score"] = min_score_value
            return self.es_client.search(index=self.es_keyword_index, body=body)

        try:
            search_result = await loop.run_in_executor(None, _search)
        except Exception as exc:
            self.logger.error(f"执行关键字检索失败: {exc}")
            raise

        hits = search_result.get("hits", {}).get("hits", []) or []
        sources: List[Dict[str, Any]] = []
        answer_fragments: List[str] = []

        for hit in hits:
            score = hit.get("_score")
            if min_score_value is not None and score is not None and score < min_score_value:
                continue

            source_doc = hit.get("_source", {})
            highlight = hit.get("highlight", {})
            highlight_fragments = highlight.get("text", [])
            preview_text = highlight_fragments[0] if highlight_fragments else source_doc.get("text", "")
            preview_text = (preview_text or "").strip()
            if len(preview_text) > 120:
                preview_text = preview_text[:120] + "..."

            sources.append(
                {
                    "filename": source_doc.get("filename", "未知文档"),
                    "content_preview": preview_text,
                    "score": score,
                }
            )
            answer_fragments.append(preview_text)

        if answer_fragments:
            answer = "\n".join(
                f"{idx + 1}. {fragment}"
                for idx, fragment in enumerate(answer_fragments)
            )
        else:
            answer = "未找到相关内容。"

        return {
            "answer": answer,
            "sources": sources,
        }

    def _format_sources_from_nodes(self, response) -> List[Dict[str, Any]]:
        """从 LlamaIndex 响应中提取来源信息"""
        sources: List[Dict[str, Any]] = []

        source_nodes = getattr(response, "source_nodes", None)
        if not source_nodes:
            return sources

        for node_with_score in source_nodes:
            metadata = getattr(node_with_score, "metadata", None)
            text = getattr(node_with_score, "text", None)
            score = getattr(node_with_score, "score", None)

            if hasattr(node_with_score, "node"):
                base_node = node_with_score.node
                metadata = getattr(base_node, "metadata", metadata) or {}
                text = getattr(base_node, "text", text)
                if text is None and hasattr(base_node, "get_content"):
                    text = base_node.get_content(metadata_mode="none")

            metadata = metadata or {}
            text = (text or "").strip()

            if len(text) > 100:
                preview = text[:100] + "..."
            else:
                preview = text

            sources.append(
                {
                    "filename": metadata.get("filename", "未知文档"),
                    "content_preview": preview,
                    "score": score,
                }
            )

        return sources
    
    def get_status(self) -> Dict[str, Any]:
        """获取系统状态信息"""
        return {
            'index_ready': self.index is not None,
            'data_dir': str(self.data_dir),
            'persist_dir': str(self.persist_dir),
            'parallel_enabled': self.enable_parallel,
            'max_workers': self.max_workers,
            'cached_files': len(self.file_cache),
            'vector_store': 'elasticsearch',
            'elasticsearch_url': self.es_url,
            'elasticsearch_index': self.es_index,
            'keyword_index': self.es_keyword_index,
            'text_analyzer': self.es_text_analyzer,
            'keyword_index_initialized': self._keyword_index_checked,
        }
    
    def __del__(self):
        """清理资源"""
        if hasattr(self, '_thread_pool'):
            self._thread_pool.shutdown(wait=False)
